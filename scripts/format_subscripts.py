#!/usr/bin/env python3
"""Conservative chemical-formula subscript auditor and formatter.

The default mode is read-only audit. Use --apply to modify only high-confidence
candidates and always write a new file.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import asdict, dataclass
from pathlib import Path

ELEMENTS = {
    "H", "He", "Li", "Be", "B", "C", "N", "O", "F", "Ne", "Na", "Mg",
    "Al", "Si", "P", "S", "Cl", "Ar", "K", "Ca", "Sc", "Ti", "V", "Cr",
    "Mn", "Fe", "Co", "Ni", "Cu", "Zn", "Ga", "Ge", "As", "Se", "Br",
    "Kr", "Rb", "Sr", "Y", "Zr", "Nb", "Mo", "Tc", "Ru", "Rh", "Pd",
    "Ag", "Cd", "In", "Sn", "Sb", "Te", "I", "Xe", "Cs", "Ba", "La",
    "Ce", "Pr", "Nd", "Pm", "Sm", "Eu", "Gd", "Tb", "Dy", "Ho", "Er",
    "Tm", "Yb", "Lu", "Hf", "Ta", "W", "Re", "Os", "Ir", "Pt", "Au",
    "Hg", "Tl", "Pb", "Bi", "Po", "At", "Rn", "Fr", "Ra", "Ac", "Th",
    "Pa", "U", "Np", "Pu", "Am", "Cm", "Bk", "Cf", "Es", "Fm", "Md",
    "No", "Lr", "Rf", "Db", "Sg", "Bh", "Hs", "Mt", "Ds", "Rg", "Cn",
    "Nh", "Fl", "Mc", "Lv", "Ts", "Og",
}
ELEMENT_ALT = "|".join(sorted(ELEMENTS, key=lambda value: (-len(value), value)))
NUMBER = r"\d+(?:\.\d+)?"
UNIT = rf"(?:{ELEMENT_ALT})(?:{NUMBER})?"
GROUP = rf"\((?:{UNIT})+\)(?:{NUMBER})?"
FORMULA_RE = re.compile(
    rf"(?<![A-Za-z0-9_.])((?:{UNIT}|{GROUP})+(?:[-–]?[xδ])?)(?![A-Za-z0-9_.])"
)
ELEMENT_RE = re.compile(ELEMENT_ALT)
SUBSCRIPT_RE = re.compile(rf"(?:{ELEMENT_ALT})({NUMBER})|\)({NUMBER})|([-–]?[xδ])$")
DIATOMIC = {"H2", "N2", "O2", "F2", "Cl2", "Br2", "I2"}
REFERENCE_HEADING_RE = re.compile(r"^\s*(参考文献|references|bibliography)\s*$", re.I)
PRE_EXCLUDE_RE = re.compile(
    r"(图|表|figure|fig\.?|table|第|章节|section|公式|式)\s*$", re.I
)
URL_RE = re.compile(r"(?:https?://|www\.|doi\s*:|10\.\d{4,9}/)", re.I)
INLINE_CODE_RE = re.compile(r"`[^`]*`")
SUBSCRIPT_UNICODE = str.maketrans("0123456789-x", "₀₁₂₃₄₅₆₇₈₉₋ₓ")


@dataclass
class Candidate:
    location: str
    text: str
    suggestion: str
    classification: str
    confidence: str
    reason: str
    start: int
    end: int
    applied: bool = False


def subscript_ranges(formula: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for match in SUBSCRIPT_RE.finditer(formula):
        for group in range(1, 4):
            if match.group(group) is not None:
                ranges.append(match.span(group))
    return ranges


def unicode_suggestion(formula: str) -> str:
    chars = list(formula)
    for start, end in reversed(subscript_ranges(formula)):
        chars[start:end] = list(formula[start:end].translate(SUBSCRIPT_UNICODE))
    return "".join(chars)


def overlaps(span: tuple[int, int], blocked: list[tuple[int, int]]) -> bool:
    return any(span[0] < end and start < span[1] for start, end in blocked)


def classify(text: str, full_text: str, start: int, end: int) -> tuple[str, str, str]:
    before = full_text[max(0, start - 20):start]
    after = full_text[end:end + 3]
    if URL_RE.search(full_text):
        return "excluded", "none", "URL/DOI context"
    if PRE_EXCLUDE_RE.search(before):
        return "excluded", "none", "figure/table/chapter/equation context"
    if after.startswith(("+", "-", "−")) and re.search(r"\d$", text):
        return "review", "low", "possible ion charge or oxidation state"
    ranges = subscript_ranges(text)
    if not ranges:
        return "unchanged", "none", "no explicit subscript character"
    element_count = len(ELEMENT_RE.findall(text))
    if element_count >= 2:
        return "chemical-formula", "high", "multiple valid element symbols"
    if text in DIATOMIC:
        return "chemical-formula", "high", "known diatomic molecule"
    if re.search(r"[-–]?[xδ]$", text):
        return "chemical-formula", "high", "explicit nonstoichiometric suffix"
    return "review", "low", "single-element token may be a sample label or variable"


def find_candidates(text: str, location: str, blocked: list[tuple[int, int]] | None = None) -> list[Candidate]:
    blocked = blocked or []
    candidates: list[Candidate] = []
    for match in FORMULA_RE.finditer(text):
        if overlaps(match.span(1), blocked):
            continue
        value = match.group(1)
        classification, confidence, reason = classify(value, text, match.start(1), match.end(1))
        if classification == "excluded":
            continue
        candidates.append(
            Candidate(
                location=location,
                text=value,
                suggestion=unicode_suggestion(value),
                classification=classification,
                confidence=confidence,
                reason=reason,
                start=match.start(1),
                end=match.end(1),
            )
        )
    return candidates


def process_text(path: Path, apply: bool, output: Path | None) -> tuple[list[Candidate], Path | None]:
    source = path.read_text(encoding="utf-8")
    lines = source.splitlines(keepends=True)
    all_candidates: list[Candidate] = []
    result: list[str] = []
    in_fence = False
    in_references = False

    for line_no, line in enumerate(lines, 1):
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
            result.append(line)
            continue
        if REFERENCE_HEADING_RE.match(stripped):
            in_references = True
            result.append(line)
            continue
        if in_fence or in_references:
            result.append(line)
            continue

        blocked = [match.span() for match in INLINE_CODE_RE.finditer(line)]
        candidates = find_candidates(line, f"line {line_no}", blocked)
        all_candidates.extend(candidates)
        if not apply:
            result.append(line)
            continue

        chars = list(line)
        for candidate in reversed(candidates):
            if candidate.confidence != "high":
                continue
            chars[candidate.start:candidate.end] = list(candidate.suggestion)
            candidate.applied = True
        result.append("".join(chars))

    if not apply:
        return all_candidates, None
    target = output or path.with_name(f"{path.stem}_subscript_fixed{path.suffix}")
    target.write_text("".join(result), encoding="utf-8")
    return all_candidates, target


def _safe_plain_run(run) -> bool:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    allowed = {f"{{{ns}}}rPr", f"{{{ns}}}t"}
    children = list(run._element)
    return (
        run._element.getparent().tag == f"{{{ns}}}p"
        and all(child.tag in allowed for child in children)
        and sum(child.tag == f"{{{ns}}}t" for child in children) == 1
    )


def _replace_run_with_segments(run, text: str, sub_positions: set[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    parent = run._element.getparent()
    index = parent.index(run._element)
    original_rpr = run._element.rPr
    segments: list[tuple[str, bool]] = []
    buffer: list[str] = []
    state: bool | None = None
    for position, char in enumerate(text):
        is_sub = position in sub_positions
        if state is not None and is_sub != state:
            segments.append(("".join(buffer), state))
            buffer = []
        state = is_sub
        buffer.append(char)
    if buffer:
        segments.append(("".join(buffer), bool(state)))

    for offset, (segment, is_sub) in enumerate(segments):
        new_run = OxmlElement("w:r")
        rpr = deepcopy(original_rpr) if original_rpr is not None else OxmlElement("w:rPr")
        if is_sub:
            vert = rpr.find(qn("w:vertAlign"))
            if vert is None:
                vert = OxmlElement("w:vertAlign")
                rpr.append(vert)
            vert.set(qn("w:val"), "subscript")
        new_run.append(rpr)
        node = OxmlElement("w:t")
        node.set(qn("xml:space"), "preserve")
        node.text = segment
        new_run.append(node)
        parent.insert(index + offset, new_run)
    parent.remove(run._element)


def _process_docx_paragraph(paragraph, location: str, apply: bool) -> list[Candidate]:
    text = "".join(run.text for run in paragraph.runs)
    candidates = find_candidates(text, location)
    if not candidates:
        return []

    offsets: list[tuple[int, int, object]] = []
    cursor = 0
    for run in paragraph.runs:
        offsets.append((cursor, cursor + len(run.text), run))
        cursor += len(run.text)

    per_run: dict[int, tuple[object, set[int]]] = {}
    for candidate in candidates:
        containing = [
            (start, end, run) for start, end, run in offsets
            if start <= candidate.start and candidate.end <= end
        ]
        if candidate.confidence != "high" or len(containing) != 1:
            if len(containing) != 1:
                candidate.classification = "review"
                candidate.confidence = "low"
                candidate.reason = "candidate crosses DOCX runs"
            continue
        start, _, run = containing[0]
        if not _safe_plain_run(run):
            candidate.classification = "review"
            candidate.confidence = "low"
            candidate.reason = "candidate is in a complex DOCX run"
            continue
        positions = per_run.setdefault(id(run), (run, set()))[1]
        for rel_start, rel_end in subscript_ranges(candidate.text):
            positions.update(range(candidate.start - start + rel_start, candidate.start - start + rel_end))
        candidate.applied = apply

    if apply:
        for run, positions in per_run.values():
            _replace_run_with_segments(run, run.text, positions)
    return candidates


def process_docx(
    path: Path, apply: bool, output: Path | None, include_tables: bool
) -> tuple[list[Candidate], Path | None]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX processing") from exc

    document = Document(path)
    candidates: list[Candidate] = []
    in_references = False
    for index, paragraph in enumerate(document.paragraphs, 1):
        if REFERENCE_HEADING_RE.match(paragraph.text.strip()):
            in_references = True
            continue
        if in_references:
            continue
        candidates.extend(_process_docx_paragraph(paragraph, f"paragraph {index}", apply))

    if include_tables:
        for table_index, table in enumerate(document.tables, 1):
            for row_index, row in enumerate(table.rows, 1):
                for cell_index, cell in enumerate(row.cells, 1):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                        location = (
                            f"table {table_index}, row {row_index}, "
                            f"cell {cell_index}, paragraph {paragraph_index}"
                        )
                        candidates.extend(_process_docx_paragraph(paragraph, location, apply))

    if not apply:
        return candidates, None
    target = output or path.with_name(f"{path.stem}_subscript_fixed{path.suffix}")
    document.save(target)
    return candidates, target


def print_report(path: Path, candidates: list[Candidate], output: Path | None) -> None:
    print(f"Source: {path}")
    print(f"Mode: {'APPLY' if output else 'AUDIT ONLY'}")
    print(f"Output: {output or '(none)'}")
    print(f"Candidates: {len(candidates)}")
    for candidate in candidates:
        status = "APPLIED" if candidate.applied else "REVIEW"
        print(
            f"[{status}] {candidate.location}: {candidate.text} -> "
            f"{candidate.suggestion} | {candidate.classification} | "
            f"{candidate.confidence} | {candidate.reason}"
        )


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit and format chemical formula subscripts conservatively")
    parser.add_argument("file", type=Path)
    parser.add_argument("--apply", action="store_true", help="write high-confidence changes to a new file")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--include-tables", action="store_true", help="include DOCX table cells")
    parser.add_argument("--json-report", type=Path)
    args = parser.parse_args()

    if not args.file.exists():
        parser.error(f"file does not exist: {args.file}")
    if args.output and not args.apply:
        parser.error("--output requires --apply")
    if args.apply and args.output and args.output.resolve() == args.file.resolve():
        parser.error("refusing to overwrite the source file")
    if args.file.suffix.lower() not in {".docx", ".md", ".txt"}:
        parser.error("supported editable formats are .docx, .md, and .txt")

    if args.apply:
        shutil.copy2(args.file, Path(f"{args.file}.bak"))

    if args.file.suffix.lower() == ".docx":
        candidates, output = process_docx(args.file, args.apply, args.output, args.include_tables)
    else:
        candidates, output = process_text(args.file, args.apply, args.output)

    print_report(args.file, candidates, output)
    if args.json_report:
        args.json_report.write_text(
            json.dumps([asdict(candidate) for candidate in candidates], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    return 0


if __name__ == "__main__":
    sys.exit(main())
