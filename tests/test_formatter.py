import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

SCRIPT = Path(__file__).parents[1] / "scripts" / "format_subscripts.py"
SPEC = importlib.util.spec_from_file_location("format_subscripts", SCRIPT)
formatter = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = formatter
SPEC.loader.exec_module(formatter)


class FormatterTests(unittest.TestCase):
    def test_common_formulas(self):
        items = formatter.find_candidates("Al2O3 Fe2O3 Fe3O4 CO2 SO2", "test")
        self.assertEqual([item.confidence for item in items], ["high"] * 5)

    def test_diatomic_and_unchanged_formula(self):
        items = formatter.find_candidates("O2 CuO", "test")
        self.assertEqual(items[0].text, "O2")
        self.assertEqual(items[0].confidence, "high")
        self.assertEqual(items[1].classification, "unchanged")

    def test_nonstoichiometric_formulas(self):
        text = "CaMnTiFeOx NOx CeO2-x O3-δ"
        items = formatter.find_candidates(text, "test")
        self.assertEqual([item.confidence for item in items], ["high"] * 4)
        self.assertEqual(items[1].suggestion, "NOₓ")

    def test_ions_are_review_only(self):
        items = formatter.find_candidates("Fe3+ Cu2+ O2−", "test")
        self.assertTrue(all(item.confidence == "low" for item in items))

    def test_numbers_and_units_are_not_candidates(self):
        text = "图2-1 表3.2 第3章 式(2-3) [2-5] 2024年 900 ℃ 2 h 50 mL 10 wt%"
        self.assertEqual(formatter.find_candidates(text, "test"), [])

    def test_reaction_coefficients_unchanged(self):
        items = formatter.find_candidates("2Fe + O2 -> 2FeO", "test")
        self.assertEqual([item.text for item in items], ["O2"])

    def test_audit_default_and_apply_output(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "sample.txt"
            source.write_text("Al2O3 and O2\n参考文献\nFe2O3", encoding="utf-8")
            candidates, output = formatter.process_text(source, False, None)
            self.assertIsNone(output)
            self.assertEqual(source.read_text(encoding="utf-8"), "Al2O3 and O2\n参考文献\nFe2O3")
            self.assertEqual(len(candidates), 2)
            _, output = formatter.process_text(source, True, None)
            self.assertEqual(output.read_text(encoding="utf-8"), "Al₂O₃ and O₂\n参考文献\nFe2O3")

    def test_docx_plain_run_applies_true_subscript(self):
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "sample.docx"
            document = Document()
            document.add_paragraph("Al2O3")
            document.save(source)
            candidates, output = formatter.process_docx(source, True, None, False)
            self.assertTrue(candidates[0].applied)
            result = Document(output)
            self.assertEqual(result.paragraphs[0].text, "Al2O3")
            values = [
                run._element.rPr.vertAlign.get(qn("w:val"))
                if run._element.rPr is not None and run._element.rPr.vertAlign is not None
                else None
                for run in result.paragraphs[0].runs
            ]
            self.assertEqual(values, [None, "subscript", None, "subscript"])

    def test_docx_cross_run_is_review_only(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "cross-run.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Al")
            paragraph.add_run("2O3")
            document.save(source)
            candidates, output = formatter.process_docx(source, True, None, False)
            self.assertEqual(candidates[0].reason, "candidate crosses DOCX runs")
            self.assertFalse(candidates[0].applied)
            self.assertEqual(Document(output).paragraphs[0].text, "Al2O3")


if __name__ == "__main__":
    unittest.main()
